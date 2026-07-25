"""Unit tests for docx_parser.densify()'s cell-text extraction — real
python-docx Table objects (not synthetic Cell lists, see test_bucket.py),
since the paragraph-boundary bug can only be reproduced via python-docx's
own paragraph/run structure.
"""
import docx

from docx_parser import densify
from normalize import ABSENT_MARKER, PRESENT_MARKER


def _one_cell_table(paragraphs):
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = paragraphs[0]
    for p in paragraphs[1:]:
        cell.add_paragraph(p)
    return table


def test_single_paragraph_cell_text_is_unchanged():
    table = _one_cell_table(["10-50 คัน"])
    rows, _ = densify(table)
    assert rows[0][0].text == "10-50 คัน"


def test_multi_paragraph_cell_joins_with_a_separator():
    """Concatenating a cell's runs with no separator at all glues
    multi-line content together with no boundary — a remark cell listing
    several tier values, one per line, becomes an unparseable run-on
    string like "10-5051-100101..." (the real A1.1 bug). Paragraphs must
    be joined with at least a space so clean_ws leaves a real boundary."""
    table = _one_cell_table(["10-50", "51-100", "101 ขึ้นไป"])
    rows, _ = densify(table)
    text = rows[0][0].text
    assert "10-50" in text and "51-100" in text and "101" in text
    assert "10-5051-100" not in text  # the bug: no boundary at all
    assert "10-50 51-100" in text


def _one_cell_table_with_run(text, font_name):
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    return table


def test_wingdings2_check_mark_translates_to_present_marker():
    """B6.2 real-data bug: a remark cell's check mark is stored as the
    plain letter 'P' under a Wingdings 2 run — the symbol font remaps it
    to a heavy checkmark glyph only at render time, so python-docx (and
    a human reading raw XML) only ever sees 'P'. Confirmed against the
    real DOCX: Wingdings 2 code point 0x50 renders as a check mark."""
    table = _one_cell_table_with_run("P", "Wingdings 2")
    rows, _ = densify(table)
    assert rows[0][0].text == PRESENT_MARKER


def test_wingdings2_x_mark_translates_to_absent_marker():
    """Same bug, the other mark: Wingdings 2 code point 0x4F ('O') renders
    as a heavy X."""
    table = _one_cell_table_with_run("O", "Wingdings 2")
    rows, _ = densify(table)
    assert rows[0][0].text == ABSENT_MARKER


def test_plain_font_letter_o_or_p_is_not_translated():
    """The translation must be font-gated — an ordinary run reading "O"
    or "P" under a normal font is just that letter, not a symbol."""
    table = _one_cell_table_with_run("P", "Angsana New")
    rows, _ = densify(table)
    assert rows[0][0].text == "P"
